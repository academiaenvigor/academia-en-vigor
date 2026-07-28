#!/usr/bin/env python3
"""Genera y maqueta las copias Word de lectura del Tema 7."""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORD_DIR = ROOT / "word"
TABLE_HELPERS = (
    Path("/root/.codex/skills/builtins/documents/scripts").resolve()
)
sys.path.insert(0, str(TABLE_HELPERS))

from table_geometry import (  # noqa: E402
    apply_table_geometry,
    column_widths_from_weights,
    section_content_width_dxa,
)


BLUE = "2E74B5"
NAVY = "17365D"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F5F8FC"
BODY = "1F1F1F"

INPUTS = {
    "parte": ROOT
    / "temas/policia-nacional/parte/"
    "tema-07-ministerio-interior-secretaria-estado-seguridad.md",
    "atestado": ROOT
    / "temas/policia-nacional/atestado/"
    "tema-07-ministerio-interior-secretaria-estado-seguridad.md",
}
OUTPUTS = {
    "parte": WORD_DIR / "Tema_07_El_Parte_Academia_En_Vigor.docx",
    "atestado": WORD_DIR / "Tema_07_El_Atestado_Academia_En_Vigor.docx",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_font(run, *, size: float | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def paragraph_style(doc, name: str):
    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name == name:
            return style
    raise KeyError(f"No existe el estilo de párrafo {name!r}")


def style_document(path: Path, view: str) -> None:
    doc = Document(path)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    section.different_first_page_header_footer = False
    doc.settings.odd_and_even_pages_header_footer = False

    normal = paragraph_style(doc, "Normal")
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BODY)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ):
        style = paragraph_style(doc, name)
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    try:
        quote = paragraph_style(doc, "Academia En Vigor · Destacado")
    except KeyError:
        quote = doc.styles.add_style(
            "Academia En Vigor · Destacado", WD_STYLE_TYPE.PARAGRAPH
        )
    quote.base_style = normal
    quote.font.name = "Calibri"
    quote.font.size = Pt(10)
    quote.font.color.rgb = RGBColor.from_string(NAVY)
    quote.paragraph_format.left_indent = Inches(0.22)
    quote.paragraph_format.right_indent = Inches(0.12)
    quote.paragraph_format.space_before = Pt(5)
    quote.paragraph_format.space_after = Pt(7)
    quote.paragraph_format.keep_together = True

    title_done = False
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Title":
            paragraph.style = paragraph_style(doc, "Heading 1")
        if paragraph.style.name == "Heading 1" and not title_done:
            title_done = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(24)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                set_font(run, size=23, color=NAVY)
                run.bold = True
        elif paragraph.style.name.startswith("Heading"):
            for run in paragraph.runs:
                set_font(run)
        elif paragraph.style.name in ("Block Text", "Quote"):
            paragraph.style = quote
            for run in paragraph.runs:
                set_font(run, size=10, color=NAVY)
        else:
            paragraph.paragraph_format.widow_control = True
            for run in paragraph.runs:
                set_font(run)

    # Segunda línea de portada.
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Temario "):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(12)
            for run in paragraph.runs:
                set_font(run, size=13, color=BLUE)
                run.bold = True
            break

    # Franja de cabecera y pie sobrios.
    for header in (
        section.header,
        section.even_page_header,
        section.first_page_header,
    ):
        hp = header.paragraphs[0]
        hp.text = "ACADEMIA EN VIGOR  ·  TEMA 7"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs:
            set_font(run, size=8.5, color=BLUE)
            run.bold = True

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = (
        f"{'EL PARTE' if view == 'parte' else 'EL ATESTADO'}"
        "  ·  Corte normativo 28/07/2026"
    )
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fp.runs:
        set_font(run, size=8, color="6D6D6D")

    content_width = section_content_width_dxa(section)
    for table in doc.tables:
        cols = len(table.columns)
        if cols == 2:
            weights = [1.25, 2.75]
        elif cols == 3:
            weights = [1.1, 1.45, 2.45]
        elif cols == 4:
            weights = [0.9, 1.25, 1.35, 1.7]
        else:
            weights = [1.0] * cols
        apply_table_geometry(
            table, column_widths_from_weights(weights, content_width)
        )
        set_repeat_table_header(table.rows[0])
        for row in table.rows:
            prevent_row_split(row)
        for cell in table.rows[0].cells:
            set_cell_shading(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_font(run, size=9, color=NAVY)
                    run.bold = True
        for row_index, row in enumerate(table.rows[1:], start=1):
            if row_index % 2 == 0:
                for cell in row.cells:
                    set_cell_shading(cell, LIGHTER_BLUE)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.05
                    paragraph.paragraph_format.widow_control = True
                    for run in paragraph.runs:
                        set_font(run, size=9)

    # Evita que una página termine en un título suelto.
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True

    # Limpieza de propiedades identificativas.
    doc.core_properties.author = "Academia En Vigor"
    doc.core_properties.last_modified_by = "Academia En Vigor"
    doc.core_properties.title = (
        "Tema 7 · El Ministerio del Interior y la Secretaría de Estado de Seguridad"
    )
    doc.core_properties.subject = (
        "Policía Nacional · Escala Básica · "
        + ("Temario Esencial" if view == "parte" else "Temario Completo")
    )
    doc.core_properties.keywords = "Tema 7, Ministerio del Interior, SES"
    doc.save(path)


def main() -> None:
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    for view, input_path in INPUTS.items():
        output_path = OUTPUTS[view]
        subprocess.run(
            [
                "pandoc",
                str(input_path),
                "--from=gfm",
                "--to=docx",
                "--standalone",
                "--output",
                str(output_path),
            ],
            check=True,
        )
        style_document(output_path, view)
        print(output_path.relative_to(ROOT))


if __name__ == "__main__":
    main()
